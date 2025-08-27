import os
import gradio as gr
from docx import Document
import datetime
from groq import Groq

# Use Groq API key stored as environment variable in Hugging Face
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

client = Groq(api_key=GROQ_API_KEY)

def generate_lesson_plan(class_, subject, topic, students, duration, date, teacher):
    prompt = f"""
You are a lesson plan expert.
Create a detailed English lesson plan using the BOPPPS model for the topic: "{topic}".
Do not repeat the basic info (class, subject, topic, teacher etc.) again in headings or body.
Include the following structured components only:
1. Bridge-In: a short, engaging paragraph.
2. 3 Clear Learning Objectives.
3. 2 Pre-Assessment Questions.
4. One detailed Participatory Learning Activity.
5. 2–3 Post-Assessment Questions.
6. Summary paragraph by the teacher.
    """

    try:
        response = client.chat.completions.create(
            model="llama3-8b-8192",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        lesson_text = response.choices[0].message.content

        doc = Document()
        doc.add_heading("BOPPPS Lesson Plan", 0)

        # Create metadata table
        table = doc.add_table(rows=0, cols=2)
        data = {
            "Class": class_,
            "Subject": subject,
            "Topic": topic,
            "No. of Students": students,
            "Duration": duration,
            "Date": date,
            "Teacher Name": teacher
        }

        for key, value in data.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = value

        doc.add_paragraph("\n")
        doc.add_paragraph(lesson_text)

        filename = f"lesson_plan_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = f"/tmp/{filename}"
        doc.save(filepath)
        return filepath

    except Exception as e:
        return f"Error: {e}"

# Build Gradio UI with form
with gr.Blocks() as demo:
    gr.Markdown("## 📘 BOPPPS Lesson Plan Generator (English Only)")

    with gr.Group():
        with gr.Row():
            class_ = gr.Textbox(label="Class", placeholder="e.g., 9th")
            subject = gr.Textbox(label="Subject", placeholder="e.g., Physics")
        with gr.Row():
            topic = gr.Textbox(label="Topic", placeholder="e.g., Newton’s First Law")
            students = gr.Textbox(label="Number of Students", placeholder="e.g., 30")
        with gr.Row():
            duration = gr.Textbox(label="Duration (in minutes)", placeholder="e.g., 40")
            date = gr.Textbox(label="Date", placeholder="e.g., 2025-05-25")
        teacher = gr.Textbox(label="Teacher Name", placeholder="e.g., Najaf Ali Sharqi")

        generate_button = gr.Button("Generate Lesson Plan")
        output_file = gr.File(label="📄 Download Word File")

    generate_button.click(
        fn=generate_lesson_plan,
        inputs=[class_, subject, topic, students, duration, date, teacher],
        outputs=output_file
    )

demo.launch()
